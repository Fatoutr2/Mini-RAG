# --- Logique du systeme RAG ---

# --- Importation des bibliotheques ---
import os # Variables d'environnements
from PyPDF2 import PdfReader # Lecture des fichiers PDF
from docx import Document # Lecture des fichiers DOCX
import faiss # Moteur de recherche vectorielle
import numpy as np # Calcul numerique
from sentence_transformers import SentenceTransformer # Embedding

# Chargement du modèle
model = SentenceTransformer("all-MiniLM-L6-v2")

chunks = []
metadata = []

# --- Chargement et chunking des documents ---
def load_documents(folder="data"):
    """
    Charge tous les fichiers du dossier :
    - TXT
    - PDF
    - DOCX
    Retourne : chunks et metadata
    """
    chunks = []
    metadata = []

    # Parcours de tous les fichiers du dossier
    for file in os.listdir(folder):
        path = os.path.join(folder, file)
        text = ""

        # ---------- TXT ----------
        if file.endswith(".txt"):
            with open(path, "r", encoding="utf-8") as f:
                text = f.read()

        # ---------- PDF ----------
        elif file.endswith(".pdf"):
            reader = PdfReader(path)
            for page in reader.pages:
                text += page.extract_text() + " "

        # ---------- DOCX ----------
        elif file.endswith(".docx"):
            doc = Document(path)

            # Texte des paragraphes
            for para in doc.paragraphs:
                text += para.text + " "

            # Texte des tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    text += row_text + " "

        else:
            continue  # ignore les autres formats

        # Découpage en chunks (~300 mots)
        words = text.split()
        for i in range(0, len(words), 300):
            chunk = " ".join(words[i:i+300])
            chunks.append(chunk)
            metadata.append({
                "source": file,
                "chunk_id": len(metadata)
            })

    return chunks, metadata

# --- CRÉATION DE L’INDEX VECTORIEL FAISS ---
def create_vector_store(chunks):
    """
    Transforme les chunks en embeddings
    et les stocke dans un index FAISS.
    """
    embeddings = model.encode(chunks) # Génération des embeddings
    dim = embeddings.shape[1] # Dimension des vecteurs
    index = faiss.IndexFlatL2(dim) # Index FAISS utilisant la distance euclidienne (L2)
    index.add(np.array(embeddings).astype("float32")) # Ajout des vecteurs dans l’index
    return index


def retrieve(query, index, chunks, metadata, embedder, top_k=5):
    """
    Recherche les chunks les plus proches de la question.

    Paramètres :
    - query    : la question de l'utilisateur
    - index    : FAISS IndexFlatL2
    - chunks   : liste des textes chunkés
    - metadata : liste des métadonnées correspondantes
    - embedder : modèle SentenceTransformer pour transformer query en vecteur
    - top_k    : nombre de résultats à retourner
    """

    # 1️⃣ Embedding de la question
    query_vec = embedder.encode([query])
    query_vec = np.array(query_vec).astype("float32")

    # 2️⃣ Recherche vectorielle FAISS
    distances, indices = index.search(query_vec, top_k)  # renvoie arrays (1, k)
    
    # 3️⃣ Construction des résultats
    results = []
    for dist, idx in zip(distances[0], indices[0]):
        if idx < len(chunks):
            results.append((chunks[idx], metadata[idx], dist))

    return results
